import os
import base64
import pdfplumber
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from docx import Document
from PIL import Image
import io

def read_file(file_path):
    _, ext = os.path.splitext(file_path)
    ext = ext.lower()

    try:
        if ext == '.txt':
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        elif ext == '.pdf':
            try:
                with pdfplumber.open(file_path) as pdf:
                    return '\n'.join(page.extract_text() for page in pdf.pages if page.extract_text())
            except Exception as e:
                return f"Error reading PDF file {file_path}: {str(e)}"
        elif ext in ['.doc', '.docx']:
            doc = Document(file_path)
            return '\n'.join([para.text for para in doc.paragraphs])
        elif ext in ['.jpg', '.jpeg', '.png', '.gif', '.bmp']:
            img = Image.open(file_path)
            return f"Image file: {file_path}\nFormat: {img.format}, Size: {img.size}, Mode: {img.mode}"        
        elif ext in ['.mp3', '.wav', '.ogg', '.m4a']:
            return f"Audio file: {file_path}"
        else:
            return f"Unsupported file format: {ext}"
    except Exception as e:
        return f"Error reading file {file_path}: {str(e)}"
    
def encode_image(image_path):
    with open(image_path, "rb") as image_file:
        return base64.b64encode(image_file.read()).decode('utf-8')

def save_output(content, output_format, file_path):
    try:
        if output_format == "PDF":
            save_as_pdf(content, file_path)
        elif output_format == "Markdown":
            with open(file_path, "w", encoding="utf-8") as f:
                f.write(content)
        elif output_format == "Document":
            doc = Document()
            doc.add_paragraph(content)
            doc.save(file_path)
        elif output_format == "Media":
            # Assuming we're saving as PNG for "Media" option
            if isinstance(content, Image.Image):
                content.save(file_path)
            else:
                raise ValueError("Cannot save non-image content as Media")
        else:
            raise ValueError(f"Unsupported output format: {output_format}")
    except Exception as e:
        raise Exception(f"Error saving output: {str(e)}")

def save_as_pdf(content, file_path):
    c = canvas.Canvas(file_path, pagesize=letter)
    width, height = letter
    c.drawString(72, height - 72, content)
    c.save()

